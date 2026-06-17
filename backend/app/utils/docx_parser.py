"""
DOCX text extraction using python-docx.
Extracts text from paragraphs and tables.
"""
import io
from docx import Document


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from a DOCX file given its bytes."""
    doc = Document(io.BytesIO(file_bytes))
    parts = []

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)
